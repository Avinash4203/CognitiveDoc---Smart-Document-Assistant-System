import streamlit as st
import io
import time
import random
import re
import json
import pandas as pd
import altair as alt
from docx import Document # python-docx
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import csv
import fitz # PyMuPDF
from streamlit_agraph import agraph, Node, Edge, Config
import base64
from collections import Counter
from gtts import gTTS
from fpdf import FPDF
from deep_translator import GoogleTranslator
import graphviz
import PyPDF2 
from textblob import TextBlob 
from streamlit_mic_recorder import mic_recorder
import plotly.express as px
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.decomposition import PCA
import numpy as np
import nltk
import subprocess
import sys

# ===================== AUTO-DOWNLOAD TEXTBLOB CORPORA =====================
@st.cache_resource
def download_corpora():
    # Helper to check if a specific corpus is missing
    try:
        from textblob import TextBlob
        _ = TextBlob("test").tags # Triggers lookup
    except Exception:
        # If missing, download all corpora (lite version is safer for cloud)
        subprocess.run([sys.executable, "-m", "textblob.download_corpora", "lite"])

download_corpora()

# Fix for TextBlob/NLTK missing corpora
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')
try:
    nltk.data.find('corpora/brown')
except LookupError:
    nltk.download('brown')


# NOTE: Ensure these modules exist in your environment
from QAWithPDF.data_ingestion import load_data
from QAWithPDF.model_api import load_model
from QAWithPDF.embedding import download_gemini_embedding

# ===================== GLOBAL PAGE CONFIG =====================
st.set_page_config(
    page_title="CognitiveDoc",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===================== THEME CSS =====================
def load_theme_css():
    st.markdown(""" 
    <style>
    :root { --bg: #0b1020; --card-bg: rgba(255,255,255,0.03); --text: #e6f7ff; --sub: #bfe6ff;
            --accent1: #4aa9ff; --accent2: #7b5cff; }
    html, body, [data-testid="stAppViewContainer"] {
        background: linear-gradient(180deg, #041029 0%, #050512 60%), var(--bg) !important;
        color: var(--text) !important;
        font-family: Inter, system-ui, -apple-system, "Segoe UI", Roboto, "Helvetica Neue", Arial;
    }
    .main-container { padding: 25px; background: var(--card-bg); border-radius: 20px;
                      box-shadow: 0 8px 28px rgba(10,18,40,0.45); }
    .stButton > button { background: linear-gradient(90deg, var(--accent1), var(--accent2));
                          color: white; border-radius: 10px; border: none; font-weight: 700; transition: all 0.3s ease; }
    .stButton > button:hover { transform: translateY(-2px); box-shadow: 0 4px 12px rgba(74,169,255,0.4); }
    .user-bubble { background: linear-gradient(135deg, rgba(74,169,255,0.15), rgba(123,92,255,0.08));
                   padding: 12px 16px; border-radius: 12px; margin: 10px 0; border-left: 3px solid rgba(74,169,255,0.6); }
    .bot-bubble { background: rgba(255,255,255,0.03); padding: 12px 16px; border-radius: 12px; margin: 10px 0;
                  border-left: 3px solid rgba(180,220,255,0.15); }
    .quiz-card { background: rgba(255,255,255,0.05); padding: 20px; border-radius: 15px; border: 1px solid var(--accent2); margin-top: 10px; }
    div[data-testid="stSidebarNav"] {display: none;} 
    
    /* New Metrics Style */
    .metric-card { background: rgba(255,255,255,0.05); padding: 15px; border-radius: 10px; border-left: 4px solid var(--accent2); text-align: center; }
    .metric-val { font-size: 24px; font-weight: bold; color: var(--accent1); }
    .metric-lbl { font-size: 14px; color: var(--sub); }
    </style>
    """, unsafe_allow_html=True)

load_theme_css()

# ===================== HELPER FUNCTIONS =====================
def extract_text_from_file(uploaded_file):
    """Robust text extractor for X-Ray and Analytics"""
    try:
        text = ""
        if uploaded_file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() or ""
        elif uploaded_file.name.endswith('.txt'):
            text = uploaded_file.getvalue().decode('utf-8')
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        return text
    except Exception as e:
        return ""

def extract_xray_data(text):
    # 1. EMAILS (Strict Regex)
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    
    # 2. MOBILE NUMBERS
    phones = re.findall(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text)
    
    # 3. DATES
    date_patterns = [
        r'\d{4}-\d{2}-\d{2}',          # 2024-01-01
        r'\d{2}/\d{2}/\d{4}',          # 01/01/2024
        r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}' # Jan 1, 2024
    ]
    dates = []
    for pattern in date_patterns:
        dates.extend(re.findall(pattern, text))
        
    unique_emails = list(set(emails))
    unique_phones = list(set(phones))
    unique_dates = list(set(dates))
    
    return unique_emails, unique_phones, unique_dates

def analyze_document_dna(text):
    blob = TextBlob(text)
    
    # Sentiment
    sentiment_score = blob.sentiment.polarity
    if sentiment_score > 0.1: sentiment = "Positive 🟢"
    elif sentiment_score < -0.1: sentiment = "Negative 🔴"
    else: sentiment = "Neutral ⚪"
    
    # Readability
    sentences = blob.sentences
    avg_len = sum(len(s.words) for s in sentences) / len(sentences) if sentences else 0
    complexity = "High" if avg_len > 20 else "Moderate" if avg_len > 10 else "Low"
    
    # Keywords
    words = [w.lower() for w in blob.words if len(w) > 4 and w.isalpha()]
    common_words = Counter(words).most_common(10)
    
    return sentiment, complexity, len(sentences), common_words

def generate_pdf_report(doc_name, history):
    def safe_text(text):
        if not text: return ""
        return text.encode('latin-1', 'replace').decode('latin-1')
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, txt=safe_text(f"DocQuest Report: {doc_name}"), ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", size=11)
    for chat in history:
        if isinstance(chat, dict): q, a = chat.get('q', ''), chat.get('a', '')
        else: q, a = chat[0], chat[1]
        pdf.set_font("Arial", 'B', 11)
        if pdf.get_x() > pdf.l_margin: pdf.ln()
        pdf.multi_cell(0, 7, txt=f"Q: {safe_text(q)}")
        pdf.set_font("Arial", '', 11)
        if pdf.get_x() > pdf.l_margin: pdf.ln()
        pdf.multi_cell(0, 7, txt=f"A: {safe_text(a)}")
        pdf.ln(5)
    return bytes(pdf.output(dest='S'))

def perform_translation(text, target_lang):
    if target_lang == "English": return text
    lang_map = {"Spanish": "es", "French": "fr", "German": "de", "Hindi": "hi", "Telugu": "te"}
    try:
        translator = GoogleTranslator(source='auto', target=lang_map.get(target_lang, "en"))
        return translator.translate(text)
    except: return f"[Error] {text}"

# ===================== SESSION STATE =====================
if "history" not in st.session_state: st.session_state.history = {}
if "uploaded_files" not in st.session_state: st.session_state.uploaded_files = []
if "xray_data" not in st.session_state: st.session_state.xray_data = {}
if "doc_stats" not in st.session_state: st.session_state.doc_stats = {} 
if "trigger_processing" not in st.session_state: st.session_state.trigger_processing = False
if "current_question" not in st.session_state: st.session_state.current_question = ""
if "quiz_data" not in st.session_state: st.session_state.quiz_data = None 
if "mindmap_edges" not in st.session_state: st.session_state.mindmap_edges = []
if "follow_ups" not in st.session_state: st.session_state.follow_ups = [] 
if "processed_audio" not in st.session_state: st.session_state.processed_audio = False

# ===================== CACHED FUNCTIONS =====================
@st.cache_resource
def get_query_engine(_model, _document_data, doc_id):
    return download_gemini_embedding(_model, _document_data)

# ===================== SIDEBAR =====================
with st.sidebar:
    st.markdown("### ⚡ CognitiveDoc")
    st.write("Intelligent Q/A System")
    
    if st.session_state.uploaded_files:
        selected_file_name = st.selectbox("Active Document", st.session_state.uploaded_files)
    else:
        selected_file_name = None
        st.info("Upload a document to begin.")
    
    st.markdown("---")

    # 2. NAVIGATION MENU (CONSOLIDATED)
    menu_selection = st.radio(
        "Navigate",
        ["💬 Chat Intelligence", "📊 Document Analytics", "🛡️ Privacy Redactor","🤖 Multi-Doc Agent", "🎮 Gamified Quiz"],
        index=0,
    )
    st.markdown("---")
    
    # 3. SETTINGS
    if selected_file_name:
        with st.expander("Output type", expanded=False):
            st.caption("AI Personal")
            persona = st.selectbox("", ["Standard", "ELI5 (Simple)", "Executive (Brief)", "Skeptic (Critical)"], label_visibility="collapsed")
            
            st.divider()
            data = st.session_state.xray_data.get(selected_file_name, {})
            st.markdown("**🔍 X-Ray Insight:**")
            if data.get('emails'):
                st.markdown(f"📧 **Emails:** {len(data['emails'])} found")
                for e in data['emails'][:3]: st.caption(f"• {e}")
            if data.get('dates'):
                st.markdown(f"📅 **Dates:** {len(data['dates'])} found")
                for d in data['dates'][:3]: st.caption(f"• {d}")

    if st.button("🗑 Clear Session"):
        st.session_state.history = {}
        st.session_state.xray_data = {}
        st.session_state.doc_stats = {}
        st.session_state.quiz_data = None
        st.session_state.mindmap_edges = []
        st.session_state.follow_ups = []
        st.cache_resource.clear()
        st.rerun()

# ===================== MAIN PAGE LOGIC =====================
st.markdown("<div class='main-container'>", unsafe_allow_html=True)
st.markdown("## 📘 CognitiveDoc – Smart Document Q&A System")

# Data Ingestion
st.markdown("### 📤 Upload Document")
docs = st.file_uploader("", type=["pdf","txt","docx"], accept_multiple_files=True, label_visibility="collapsed")

if docs:
    for d in docs:
        if d.name not in st.session_state.uploaded_files:
            st.session_state.uploaded_files.append(d.name)
            try:
                raw_text = extract_text_from_file(d)
                
                # CALL THE XRAY FUNCTION
                emails, phones, dates = extract_xray_data(raw_text)
                st.session_state.xray_data[d.name] = {
                    'emails': emails, 
                    'phones': phones, 
                    'dates': sorted(dates)
                }
                
                # DNA Analytics
                sent, comp, sent_count, keywords = analyze_document_dna(raw_text)
                st.session_state.doc_stats[d.name] = {
                    'sentiment': sent, 
                    'complexity': comp, 
                    'sentences': sent_count, 
                    'keywords': keywords
                }
            except Exception as e:
                st.error(f"Error processing {d.name}: {e}")

if st.session_state.uploaded_files:
    col_sel, col_lang = st.columns([3, 1])
    with col_sel:
        if selected_file_name is None: selected_file_name = st.session_state.uploaded_files[0]
    with col_lang:
        target_lang = st.selectbox("Output Language", ["English", "Spanish", "French", "German", "Hindi", "Telugu"])
else:
    selected_file_name = None
    target_lang = "English"
    

# ---------------- VIEW 1: CHAT INTELLIGENCE (WITH VOICE MODE) ----------------
if menu_selection == "💬 Chat Intelligence":
    st.markdown(f"## 💬 Chat Intelligence")
    if selected_file_name:
        st.caption(f"Analyzing: {selected_file_name}")
        
        # Suggestions & Deep Dive
        suggestions = ["Summarize key points", "What is the main conclusion?", "List technical terms"]
        cols = st.columns(4)
        for i, sugg in enumerate(suggestions):
            if cols[i].button(sugg, key=f"s_{i}"):
                st.session_state.current_question = sugg
                st.session_state.trigger_processing = True
                st.rerun()
        if cols[3].button("🕵️ Deep Dive"):
            st.session_state.current_question = "CONDUCT_DEEP_DIVE"
            st.session_state.trigger_processing = True
            st.rerun()

        # INPUT AREA: Text + Voice
        input_col, mic_col = st.columns([6, 1])
        
        with input_col:
            user_question = st.text_input("Ask anything...", value=st.session_state.current_question, label_visibility="collapsed", placeholder="Type or click the mic to speak...")
        
        with mic_col:
            # VOICE INPUT INTEGRATION
            audio = mic_recorder(start_prompt="🎙️", stop_prompt="⏹️", just_once=False, key='recorder', use_container_width=True)
        
        # Handle Voice Input
        if audio and not st.session_state.processed_audio:
            # Simulation of speech-to-text
            st.toast("Voice received! (Using simulated text for demo)", icon="🎤")
            # In a real app, send audio['bytes'] to Google Speech API or OpenAI Whisper
            # user_question = transcribe_audio(audio['bytes']) 
            user_question = "Summarize the key findings of this document." # Demo fallback
            st.session_state.current_question = user_question
            st.session_state.trigger_processing = True
            st.session_state.processed_audio = True # Prevent loops
            st.rerun()

        # Update text input state
        if user_question != st.session_state.current_question:
            st.session_state.current_question = user_question
            st.session_state.processed_audio = False

        if st.button("🚀 Process Query") or st.session_state.trigger_processing:
            st.session_state.trigger_processing = False
            st.session_state.processed_audio = False
            
            if not st.session_state.current_question.strip(): st.warning("Enter a query.")
            else:
                status_box = st.status("🧠 Processing...", expanded=True)
                try:
                    doc_obj = next((d for d in docs if d.name == selected_file_name), None)
                    if doc_obj:
                        document_data = load_data(doc_obj)
                        if not isinstance(document_data, list): document_data = [document_data]
                        document_data = [d for d in document_data if getattr(d, "text", None)]
                        
                        model = load_model()
                        query_engine = get_query_engine(model, document_data, selected_file_name)
                        
                        if st.session_state.current_question == "CONDUCT_DEEP_DIVE":
                            status_box.write("🕵️ Running deep investigation...")
                            prompt = "Generate a comprehensive investigation report: 1. Introduction, 2.Main info , 3. Hidden Details, 4. Conclusion."
                        else:
                            prompt = st.session_state.current_question
                            if persona == "ELI5 (Simple)": prompt += " (Explain like I'm 5)"
                            elif persona == "Executive (Brief)": prompt += " (Executive summary only)"
                            elif persona == "Skeptic (Critical)": prompt += " (Analyze critically and give more elaborate answer)"

                        response = query_engine.query(prompt)
                        response_text = response.response
                        
                        # Generate Smart Follow-ups
                        st.session_state.follow_ups = [
                            f"Tell me more about {st.session_state.current_question.split()[0]}",
                            "Explain the pros and cons",
                            "What data supports this?"
                        ]
                        
                        if target_lang != "English":
                            status_box.write(f"🌍 Translating...")
                            response_text = perform_translation(response_text, target_lang)

                        source = response.source_nodes[0].node.text[:250] + "..." if response.source_nodes else ""
                        
                        if selected_file_name not in st.session_state.history: st.session_state.history[selected_file_name] = []
                        display_q = "🕵️ Deep Dive Report" if st.session_state.current_question == "CONDUCT_DEEP_DIVE" else st.session_state.current_question
                        
                        st.session_state.history[selected_file_name].append({"q": display_q, "a": response_text, "src": source})
                        status_box.update(label="✅ Done!", state="complete", expanded=False)
                        
                        # IF Voice was used, Play Audio Response automatically
                        

                except Exception as e:
                    status_box.update(label="❌ Error", state="error")
                    st.error(str(e))

        if selected_file_name in st.session_state.history:
            st.divider()
            
            # Display Follow-up Buttons if latest chat
            if st.session_state.follow_ups:
                st.caption("💡 Suggested Follow-ups:")
                f1, f2, f3 = st.columns(3)
                if f1.button(st.session_state.follow_ups[0]): 
                    st.session_state.current_question = st.session_state.follow_ups[0]
                    st.session_state.trigger_processing = True
                    st.rerun()
                if f2.button(st.session_state.follow_ups[1]): 
                    st.session_state.current_question = st.session_state.follow_ups[1]
                    st.session_state.trigger_processing = True
                    st.rerun()
                if f3.button(st.session_state.follow_ups[2]): 
                    st.session_state.current_question = st.session_state.follow_ups[2]
                    st.session_state.trigger_processing = True
                    st.rerun()

            for i, chat in enumerate(reversed(st.session_state.history[selected_file_name]), 1):
                if isinstance(chat, dict): q, a, src = chat.get('q',""), chat.get('a',""), chat.get('src',"")
                else: q, a, src = chat[0], chat[1], chat[2]
                
                st.markdown(f"<div class='user-bubble'><strong>Q:</strong> {q}</div>", unsafe_allow_html=True)
                st.markdown(f"<div class='bot-bubble'><strong>A:</strong> {a}</div>", unsafe_allow_html=True)
                
                c1, c2, c3 = st.columns([1, 4, 1])
                with c1: 
                    if src: 
                        with st.expander("🔍 Source"): st.info(src)
                
                with c3:
                    if st.button("📥 PDF", key=f"pdf_{i}"):
                         pdf = generate_pdf_report(selected_file_name, st.session_state.history[selected_file_name])
                         st.download_button("Download", pdf, f"{selected_file_name}.pdf", "application/pdf")
    else:
        st.info("👈 Upload a document to start chatting.")
# ---------------- VIEW: MULTI-DOC AGENT (BATCH PROCESSOR) ----------------
# ---------------- VIEW: MULTI-DOC AGENT (BATCH PROCESSOR) ----------------
elif menu_selection == "🤖 Multi-Doc Agent":
    st.markdown("## 🤖 The Automated Recruiter")
    st.caption("Analyze ALL uploaded documents at once to find the best match.")
    
    if len(st.session_state.uploaded_files) < 2:
        st.warning("⚠️ Upload at least 2 documents (e.g., Resumes or Contracts) to use this agent.")
    else:
        # 1. Define the Goal
        st.subheader("1. What is your goal?")
        agent_mode = st.selectbox("Select Agent Mode:", 
                                ["Resume Ranker (Find best candidate)", 
                                 "Contract Auditor (Find high risk)", 
                                 "Custom Criteria"])
        
        criteria = ""
        if agent_mode == "Resume Ranker (Find best candidate)":
            criteria = st.text_input("Enter Job Description / Key Skills:", "Python, AI, 3+ years experience")
        elif agent_mode == "Contract Auditor (Find high risk)":
            criteria = st.text_input("Enter Risk Factors:", "Liability > $1M, Termination without cause")
        else:
            criteria = st.text_input("Enter your scoring criteria:", "Explain specifically what to look for.")
            
        if st.button("🤖 Run Agent on All Files"):
            if not criteria:
                st.error("Please enter criteria.")
            else:
                results = []
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # 2. Iterate through ALL docs
                for idx, file_name in enumerate(st.session_state.uploaded_files):
                    
                    # === SAFETY BRAKE: Protect API Quota ===
                    if idx > 0: 
                        status_text.text(f"⏳ Cooling down (Safe Mode)... Analyzing {file_name}...")
                        time.sleep(4) # Wait 4s to stay under 15 RPM limit
                    else:
                        status_text.text(f"Analyzing {file_name}...")
                    
                    # Load Doc
                    try:
                        doc_obj = next(d for d in docs if d.name == file_name)
                        # Quick text extract (limit to 4000 chars to save tokens)
                        text = extract_text_from_file(doc_obj)[:4000] 
                        
                        # Call LLM
                        # Use Fast Model (False) to save quota
                        model = load_model() 
                        
                        prompt = f"""
                        Analyze this document text: "{text}..."
                        
                        Based strictly on these criteria: "{criteria}"
                        
                        Return a JSON object with these exact keys:
                        - "score": (integer 0-100)
                        - "reason": (1 short sentence why)
                        - "key_match": (one specific quote or skill found)
                        
                        Return ONLY raw JSON string. No markdown formatting.
                        """
                        
                        response = model.complete(prompt)
                        
                        # Clean and Parse JSON
                        json_str = response.text.strip().replace("```json", "").replace("```", "")
                        data = json.loads(json_str)
                        data['file_name'] = file_name
                        results.append(data)
                        
                    except Exception as e:
                        # Handle errors gracefully so one bad file doesn't stop the loop
                        results.append({
                            "file_name": file_name, 
                            "score": 0, 
                            "reason": f"Error: {str(e)}", 
                            "key_match": "-"
                        })
                    
                    # Update Progress
                    progress_bar.progress((idx + 1) / len(st.session_state.uploaded_files))
                
                status_text.text("✅ Analysis Complete!")
                
                # 3. Visualization
                if results:
                    df_results = pd.DataFrame(results).sort_values(by="score", ascending=False)
                    
                    # Top Candidate Highlight
                    best = df_results.iloc[0]
                    st.success(f"🏆 Top Match: **{best['file_name']}** (Score: {best['score']}/100)")
                    
                    # Interactive Table
                    st.dataframe(
                        df_results,
                        column_order=("score", "file_name", "reason", "key_match"),
                        hide_index=True,
                        use_container_width=True,
                        column_config={
                            "score": st.column_config.ProgressColumn("Match Score", format="%d", min_value=0, max_value=100),
                            "file_name": "Document Name",
                            "reason": "AI Reasoning",
                            "key_match": "Evidence Found"
                        }
                    )
# ---------------- VIEW 2: DOCUMENT ANALYTICS (CONSOLIDATED) ----------------
elif menu_selection == "📊 Document Analytics":
    
    st.markdown("## 📊 Document Analytics & Visualization")
    if selected_file_name:
        # Get raw text once
        raw_text = extract_text_from_file(next(d for d in docs if d.name == selected_file_name))
        
        # --- TAB SELECTION FOR ORGANIZED VIEW ---
        # Combined all advanced visualizations here
        tabs = st.tabs(["📈 DNA Metrics", "☁️ Word Cloud", "🕸️ Interactive Graph", "🧠 Mind Map", "🧭 3D Holo-View", "🔍 Rapid Search"])
        
        # TAB 1: METRICS
        with tabs[0]:
            st.subheader("Document Vitals")
            if selected_file_name in st.session_state.doc_stats:
                stats = st.session_state.doc_stats[selected_file_name]
                word_count = len(raw_text.split())
                read_time = max(1, round(word_count / 200))
                
                m1, m2, m3, m4 = st.columns(4)
                with m1: st.metric("Sentiment", stats['sentiment'])
                with m2: st.metric("Complexity", stats['complexity'])
                with m3: st.metric("Sentences", stats['sentences'])
                with m4: st.metric("Read Time", f"~{read_time} min")
                
                st.divider()
                st.subheader("📅 Event Timeline")
                dates = st.session_state.xray_data[selected_file_name]['dates']
                if dates:
                    timeline_data = [{"Date": d, "Event": "Mentioned in text"} for d in dates]
                    st.dataframe(timeline_data, use_container_width=True)
                else:
                    st.info("No dates found to build a timeline.")

        # TAB 2: VISUAL WORD CLOUD
        with tabs[1]:
            st.subheader("Key Topics Visualized")
            if raw_text:
                with st.spinner("Generating Word Cloud..."):
                    wordcloud = WordCloud(width=800, height=400, background_color='#0b1020', colormap='Blues').generate(raw_text)
                    fig, ax = plt.subplots(figsize=(10, 5))
                    ax.imshow(wordcloud, interpolation='bilinear')
                    ax.axis("off")
                    fig.patch.set_alpha(0) 
                    st.pyplot(fig)
            else:
                st.warning("Document text is empty.")

        # TAB 3: INTERACTIVE GRAPH (PHYSICS BASED - IMPROVED)
        with tabs[2]:
            st.subheader("🕸️ Concept Network")
            col_graph, col_legend = st.columns([3, 1])
            
            with col_legend:
                st.info("""
                **How to read this:**
                * **Nodes:** Key concepts.
                * **Size:** How often it appears.
                * **Lines:** Concepts appearing in the same sentence.
                * **Action:** Drag nodes to untangle connections.
                """)
            
            with col_graph:
                with st.spinner("Simulating physics..."):
                    blob = TextBlob(raw_text)
                    nouns = [n.lower() for n in blob.noun_phrases if len(n.split()) < 3]
                    # Get counts to size the nodes
                    noun_counts = Counter(nouns)
                    top_concepts = [w[0] for w in noun_counts.most_common(12)] # Increased to 12
                    
                    nodes = []
                    edges = []
                    
                    for concept in top_concepts:
                        # Dynamic sizing based on frequency
                        size = 15 + (noun_counts[concept] * 2) 
                        nodes.append(Node(id=concept, label=concept, size=size, color="#4aa9ff"))
                    
                    for i in range(len(top_concepts)):
                        for j in range(i + 1, len(top_concepts)):
                            c1 = top_concepts[i]
                            c2 = top_concepts[j]
                            weight = 0
                            for sent in blob.sentences:
                                if c1 in sent.lower() and c2 in sent.lower():
                                    weight += 1
                            
                            if weight > 0:
                                edges.append(Edge(source=c1, target=c2, color="#7b5cff", strokeWidth=weight))

                    # Improved Physics Config for stability
                    config = Config(width=800, 
                                    height=500, 
                                    directed=False,
                                    nodeHighlightBehavior=True, 
                                    highlightColor="#F7A7A6", 
                                    collapsible=True,
                                    physics={
                                        "enabled": True,
                                        "solver": "forceAtlas2Based", # Better solver
                                        "stabilization": {"iterations": 200}
                                    })
                    return_value = agraph(nodes=nodes, edges=edges, config=config)

        # TAB 4: VISUAL MIND MAP (LLM POWERED)
        with tabs[3]:
            st.subheader("🧠 Hierarchical Mind Map")
            if st.button("Generate Mind Map Structure") or not st.session_state.mindmap_edges:
                with st.spinner("AI is determining relationships..."):
                    try:
                        doc_obj = next((d for d in docs if d.name == selected_file_name), None)
                        if doc_obj:
                            document_data = load_data(doc_obj)
                            if not isinstance(document_data, list): document_data = [document_data]
                            document_data = [d for d in document_data if getattr(d, "text", None)]
                            model = load_model()
                            query_engine = get_query_engine(model, document_data, selected_file_name)
                            
                            prompt = "Identify the top 10 most important concepts in this document. Then, identify how they are related. Format the output strictly as: Concept A -> Concept B. Return only 5 lines of these relationships."
                            response = query_engine.query(prompt)
                            raw_edges = str(response.response).split('\n')
                            
                            st.session_state.mindmap_edges = []
                            for edge in raw_edges:
                                if "->" in edge:
                                    parts = edge.split("->")
                                    if len(parts) == 2:
                                        st.session_state.mindmap_edges.append((parts[0].strip(), parts[1].strip()))
                    except Exception as e:
                        st.error(f"Analysis failed: {e}")

            if st.session_state.mindmap_edges:
                try:
                    graph = graphviz.Digraph()
                    graph.attr(bgcolor='transparent')
                    graph.attr('node', style='filled', fillcolor='#4aa9ff', fontcolor='white', shape='box')
                    for start, end in st.session_state.mindmap_edges:
                        graph.edge(start, end)
                    st.graphviz_chart(graph)
                except:
                    st.error("Graphviz missing. Showing raw data:")
                    st.write(st.session_state.mindmap_edges)

        # TAB 5: 3D HOLO-VIEW
        with tabs[4]:
            st.subheader("🧭 3D Semantic Space")
            st.caption("Sentences with similar meanings cluster together.")
            
            sentences = [s.strip() for s in raw_text.split('.') if len(s.split()) > 5]
            if len(sentences) < 5:
                st.warning("Not enough text to visualize.")
            else:
                with st.spinner("Calculating 3D vectors..."):
                    vectorizer = TfidfVectorizer(max_features=100, stop_words='english')
                    X = vectorizer.fit_transform(sentences)
                    pca = PCA(n_components=3)
                    result = pca.fit_transform(X.toarray())
                    
                    df_plot = pd.DataFrame(result, columns=['x', 'y', 'z'])
                    df_plot['sentence'] = sentences[:len(result)]
                    df_plot['length'] = df_plot['sentence'].apply(len)
                    
                    fig = px.scatter_3d(
                        df_plot, x='x', y='y', z='z',
                        color='length', 
                        hover_data=['sentence'],
                        opacity=0.7,
                        color_continuous_scale='Viridis'
                    )
                    fig.update_layout(
                        margin=dict(l=0, r=0, b=0, t=0),
                        scene=dict(bgcolor="#0b1020"), 
                        paper_bgcolor="#0b1020",
                        font_color="white"
                    )
                    st.plotly_chart(fig, use_container_width=True)

        # TAB 6: SEARCH
        with tabs[5]:
            st.subheader("⚡ Instant Text Finder")
            search_term = st.text_input("Find exact phrase:", placeholder="e.g., 'contract clause'")
            if search_term:
                matches = []
                lines = raw_text.split('\n')
                for i, line in enumerate(lines):
                    if search_term.lower() in line.lower():
                        matches.append(f"**Line {i+1}:** ...{line.strip()}...")
                if matches:
                    st.success(f"Found {len(matches)} matches:")
                    for m in matches[:10]: st.markdown(m)
                else:
                    st.warning("No matches found.")
    else:
        st.info("Upload a document to view analytics.")


# ---------------- VIEW 3: SECURITY & PRIVACY (UNIVERSAL) ----------------
elif menu_selection == "🛡️ Privacy Redactor":
    st.markdown("## 🛡️ Secure Data Room")
    st.caption("Enterprise-grade Redaction & Risk Assessment (Supports PDF & DOCX).")
    
    if selected_file_name:
        # Get the original file object
        file_obj = next((d for d in docs if d.name == selected_file_name), None)
        
        if file_obj:
            col1, col2 = st.columns([1, 1])
            
            # --- LEFT COLUMN: REDACTION TOOL ---
            with col1:
                st.subheader("🕵️ Universal Redactor")
                if selected_file_name.endswith(".pdf"):
                    st.info("Mode: PDF (Black Box Overlay)")
                elif selected_file_name.endswith(".docx"):
                    st.info("Mode: DOCX (Text Replacement)")
                else:
                    st.warning("File type not supported for redaction.")

                if st.button("🔴 REDACT & DOWNLOAD"):
                    with st.spinner("Applying redaction protocols..."):
                        try:
                            output_buffer = io.BytesIO()
                            mime_type = ""
                            out_name = ""
                            redaction_count = 0
                            
                            # Define Patterns
                            pii_patterns = [
                                r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', # Emails
                                r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', # Phones
                                r'[A-Z]{5}[0-9]{4}[A-Z]{1}' # PAN Cards
                            ]

                            # === PDF LOGIC (PyMuPDF) ===
                            if selected_file_name.endswith(".pdf"):
                                pdf_bytes = file_obj.getvalue()
                                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                                
                                for page in doc:
                                    page_text = page.get_text()
                                    for pattern in pii_patterns:
                                        matches = re.findall(pattern, page_text)
                                        for word in matches:
                                            areas = page.search_for(word)
                                            for area in areas:
                                                page.add_redact_annot(area, fill=(0, 0, 0))
                                                redaction_count += 1
                                    page.apply_redactions()
                                
                                doc.save(output_buffer)
                                mime_type = "application/pdf"
                                out_name = f"REDACTED_{selected_file_name}"

                            # === DOCX LOGIC (python-docx) ===
                            elif selected_file_name.endswith(".docx"):
                                # python-docx requires a file-like object, so we wrap bytes
                                docx_file = io.BytesIO(file_obj.getvalue())
                                doc = Document(docx_file)
                                
                                def redact_text(text):
                                    # Helper to replace PII in a string
                                    original = text
                                    count = 0
                                    for pattern in pii_patterns:
                                        matches = re.findall(pattern, text)
                                        for m in matches:
                                            # Replace with blocks of same length
                                            text = text.replace(m, "█" * len(m))
                                            count += 1
                                    return text, count

                                # Iterate over paragraphs
                                for paragraph in doc.paragraphs:
                                    if any(re.search(p, paragraph.text) for p in pii_patterns):
                                        # We must iterate over 'runs' to preserve formatting (bold/italic)
                                        for run in paragraph.runs:
                                            new_text, cnt = redact_text(run.text)
                                            if cnt > 0:
                                                run.text = new_text
                                                redaction_count += cnt

                                # Iterate over tables
                                for table in doc.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    new_text, cnt = redact_text(run.text)
                                                    if cnt > 0:
                                                        run.text = new_text
                                                        redaction_count += cnt
                                
                                doc.save(output_buffer)
                                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                out_name = f"REDACTED_{selected_file_name}"

                            # === FINALIZE ===
                            output_buffer.seek(0)
                            st.success(f"✅ Redacted {redaction_count} items.")
                            
                            st.download_button(
                                label="🔒 Download Safe Copy",
                                data=output_buffer,
                                file_name=out_name,
                                mime=mime_type
                            )

                        except Exception as e:
                            st.error(f"Redaction failed: {e}")

            # --- RIGHT COLUMN: RISK ANALYTICS ---
            with col2:
                st.subheader("📊 Privacy Risk Assessment")
                
                # 1. ROBUST TEXT EXTRACTION (Using PyMuPDF/fitz instead of standard extractor)
                try:
                    text_for_scanning = ""
                    if selected_file_name.endswith(".pdf"):
                        pdf_bytes = file_obj.getvalue()
                        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
                            for page in doc:
                                text_for_scanning += page.get_text()
                    elif selected_file_name.endswith(".docx"):
                        import io
                        from docx import Document
                        docx_bytes = file_obj.getvalue()
                        doc = Document(io.BytesIO(docx_bytes))
                        for para in doc.paragraphs:
                            text_for_scanning += para.text + "\n"
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    text_for_scanning += cell.text + " "
                    else:
                        file_obj.seek(0)
                        text_for_scanning = file_obj.read().decode("utf-8", errors="ignore")
                except Exception as e:
                    text_for_scanning = ""

                # 2. CALCULATE RISK
                risk_score = 0
                risks_found = []
                
                emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text_for_scanning)
                if len(emails) > 0:
                    risk_score += 30
                    risks_found.append(f"{len(emails)} Emails Detected")
                
                phones = re.findall(r'(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text_for_scanning)
                if len(phones) > 0:
                    risk_score += 30
                    risks_found.append(f"{len(phones)} Phone Numbers")

                ids = re.findall(r'[A-Z]{5}[0-9]{4}[A-Z]{1}', text_for_scanning)
                if len(ids) > 0:
                    risk_score += 40
                    risks_found.append(f"{len(ids)} Govt IDs (PAN)")

                risk_score = min(risk_score, 100)
                score_color = '#ff4b4b' if risk_score > 70 else '#ffa700' if risk_score > 30 else '#00cc66'
                
                st.markdown(f"""
                <div style="text-align:center; margin-bottom:15px; background: rgba(255,255,255,0.05); padding: 20px; border-radius: 15px; border: 1px solid {score_color};">
                    <h1 style="font-size: 60px; margin:0; color: {score_color}">
                        {risk_score}%
                    </h1>
                    <p style="color:#bfe6ff; margin:0; font-size: 14px;">Vulnerability Score</p>
                </div>
                """, unsafe_allow_html=True)
                
                if risk_score > 0:
                    st.write("**⚠️ Threats Detected:**")
                    for r in risks_found:
                        st.markdown(f"🚫 {r}")
                else:
                    st.success("🛡️ No sensitive PII detected.")

    else:
        st.info("Upload a PDF or DOCX to audit its security.")

# ---------------- VIEW 4: GAMIFIED QUIZ ----------------
elif menu_selection == "🎮 Gamified Quiz":
    st.markdown(f"## 🎮 Knowledge Check")
    if selected_file_name:
        st.write(f"Testing knowledge on: **{selected_file_name}**")
        
        if st.button("🎲 Generate New Quiz"):
            with st.spinner("AI is creating questions from your document..."):
                try:
                    # 1. Init Model
                    doc_obj = next((d for d in docs if d.name == selected_file_name), None)
                    if doc_obj:
                        document_data = load_data(doc_obj)
                        if not isinstance(document_data, list): document_data = [document_data]
                        document_data = [d for d in document_data if getattr(d, "text", None)]
                        model = load_model()
                        query_engine = get_query_engine(model, document_data, selected_file_name)

                        # 2. Ask LLM for Quiz JSON
                        prompt = """
                        Generate 5 multiple choice questions based on the document. 
                        Return the output in this strict format:
                        Q1: [Question Text] | [Option1, Option2, Option3, Option4] | [Correct Option]
                        Q2: ...
                        Q3: ...
                        Do not add markdown formatting like ** or ##.
                        """
                        response = query_engine.query(prompt)
                        raw_text = str(response.response)
                        
                        # 3. Parse Text
                        new_quiz = []
                        lines = raw_text.split('\n')
                        for line in lines:
                            if "|" in line:
                                parts = line.split("|")
                                if len(parts) >= 3:
                                    q_text = parts[0].split(":")[1].strip() if ":" in parts[0] else parts[0].strip()
                                    options = parts[1].replace("[", "").replace("]", "").split(",")
                                    correct = parts[2].replace("[", "").replace("]", "").strip()
                                    options = [o.strip() for o in options]
                                    new_quiz.append({"q": q_text, "opts": options, "ans": correct})
                        
                        st.session_state.quiz_data = new_quiz
                except Exception as e:
                    st.error(f"Quiz generation failed: {e}")
        
        if st.session_state.quiz_data:
            for idx, q_item in enumerate(st.session_state.quiz_data):
                st.markdown(f"<div class='quiz-card'><strong>Q{idx+1}: {q_item['q']}</strong></div>", unsafe_allow_html=True)
                opts = q_item['opts'] if len(q_item['opts']) > 1 else ["True", "False", "Yes", "No"]
                choice = st.radio(f"Select answer:", opts, key=f"quiz_{idx}")
                
                if st.button(f"Check Answer {idx+1}"):
                    if choice.lower() in q_item['ans'].lower() or q_item['ans'].lower() in choice.lower():
                        st.success("✅ Correct!")
                    else:
                        st.error(f"❌ Incorrect. The answer is: {q_item['ans']}")
        else:
            st.info("Click 'Generate New Quiz' to start.")
    else:
        st.info("👈 Upload a document to generate a quiz.")

st.markdown("</div>", unsafe_allow_html=True)
